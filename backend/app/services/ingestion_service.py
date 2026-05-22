import os
import pypdf
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.logger import logger
from app.repositories.document_repository import DocumentRepository
from app.services.embedding_service import EmbeddingService


class IngestionService:
    def __init__(self, db):
        self.repo = DocumentRepository(db)
        self.embedding_service = EmbeddingService()

    def extract_text(self, file_path: str):
        ext = os.path.splitext(file_path)[1].lower()

        logger.info(f"Starting text extraction | file={file_path} | ext={ext}")

        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            logger.info(f"TXT extracted | characters={len(text)}")
            return text

        if ext == ".pdf":
            reader = pypdf.PdfReader(file_path)
            text = ""

            logger.info(f"PDF pages detected | pages={len(reader.pages)}")

            for page_number, page in enumerate(reader.pages, start=1):
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

                logger.info(
                    f"PDF page processed | page={page_number} | extracted={bool(extracted)}"
                )

            logger.info(f"PDF extracted | characters={len(text)}")
            return text

        if ext == ".docx":
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            logger.info(
                f"DOCX extracted | paragraphs={len(doc.paragraphs)} | characters={len(text)}"
            )
            return text

        logger.error(f"Unsupported file type | ext={ext}")
        raise Exception(f"Unsupported file type: {ext}")

    def ingest_uploaded_file(self, file_path: str, title: str):
        logger.info(f"Document ingestion started | title={title} | path={file_path}")

        content = self.extract_text(file_path)

        if not content.strip():
            logger.error(f"No readable content found | title={title}")
            raise Exception("No readable content found in document")

        logger.info(f"Document content ready | title={title} | characters={len(content)}")

        document = self.repo.create_document(
            title=title,
            content=content,
            source_type="uploaded"
        )

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50
        )

        chunks = splitter.split_text(content)

        logger.info(
            f"Document chunking completed | document_id={document.id} | chunks={len(chunks)}"
        )

        for index, chunk in enumerate(chunks, start=1):
            logger.info(
                f"Embedding chunk started | document_id={document.id} | chunk={index}/{len(chunks)} | chars={len(chunk)}"
            )

            embedding = self.embedding_service.embed_text(chunk)

            self.repo.create_chunk(
                document_id=document.id,
                chunk_text=chunk,
                embedding=embedding,
                metadata={
                    "source": title,
                    "chunk_index": index
                }
            )

            logger.info(
                f"Chunk stored | document_id={document.id} | chunk={index}/{len(chunks)}"
            )

        logger.info(
            f"Document ingestion completed | document_id={document.id} | chunks={len(chunks)}"
        )

        return {
            "message": "Document uploaded and indexed successfully",
            "document_id": document.id,
            "chunks_created": len(chunks)
        }